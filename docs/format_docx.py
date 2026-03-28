"""
Recrée DocumentationIOT.docx avec une mise en forme professionnelle
Lancer : python format_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── PALETTE ─────────────────────────────────────────────────────
NAVY      = RGBColor(0x0B, 0x1F, 0x4B)
BLUE      = RGBColor(0x1A, 0x56, 0xBB)
ACCENT    = RGBColor(0x00, 0x8C, 0xD7)
DARK      = RGBColor(0x22, 0x22, 0x22)
GRAY      = RGBColor(0x55, 0x55, 0x55)
MID_GRAY  = RGBColor(0x88, 0x88, 0x88)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY   = "0B1F4B"
HEX_BLUE   = "1A56BB"
HEX_ACCENT = "008CD7"
HEX_EVEN   = "EBF0FA"
HEX_ODD    = "F7F9FC"
HEX_TOTAL  = "C2D3F0"

# ─── HELPERS XML ─────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_pad(cell, top=80, bot=80, left=120, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bot),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    bdr = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'),'none'); bdr.append(b)
    for old in tblPr.findall(qn('w:tblBorders')): tblPr.remove(old)
    tblPr.append(bdr)

def add_hr(doc, color=HEX_ACCENT, sz=8):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz))
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

def sp(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def add_page_number(para):
    run = para.add_run()
    for tag, text in [('begin', None), (None, 'PAGE'), ('end', None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = text; run._r.append(it)

# ─── BLOCS SÉMANTIQUES ───────────────────────────────────────────

def section_title(doc, number, title):
    p = doc.add_paragraph()
    sp(p, before=20, after=4)
    rn = p.add_run(f"{number}.  ")
    rn.bold = True; rn.font.size = Pt(15); rn.font.color.rgb = ACCENT
    rt = p.add_run(title)
    rt.bold = True; rt.font.size = Pt(15); rt.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
    bot.set(qn('w:space'),'2'); bot.set(qn('w:color'), HEX_BLUE)
    pBdr.append(bot); pPr.append(pBdr)

def sub_title(doc, text):
    p = doc.add_paragraph()
    sp(p, before=10, after=3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = BLUE

def body(doc, text, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    sp(p, before=2, after=5)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.italic = italic; r.bold = bold
    r.font.color.rgb = color or DARK

def bullet(doc, text):
    p = doc.add_paragraph()
    sp(p, before=1, after=2)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    r = p.add_run("▸  " + text)
    r.font.size = Pt(10.5); r.font.color.rgb = DARK

def big_val(doc, label, value, sub_text=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=6, after=2)
    if label:
        rl = p.add_run(label + "\n")
        rl.bold = True; rl.font.size = Pt(11); rl.font.color.rgb = GRAY
    rv = p.add_run(value)
    rv.bold = True; rv.font.size = Pt(28); rv.font.color.rgb = NAVY
    if sub_text:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp(p2, before=0, after=10)
        r2 = p2.add_run(sub_text)
        r2.font.size = Pt(11); r2.font.color.rgb = BLUE
    doc.add_paragraph()

def table(doc, headers, rows, col_widths=None, center_from=1):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)

    # header row
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        set_cell_bg(c, HEX_NAVY); set_cell_pad(c)
        p = c.paragraphs[0]; p.clear()
        r = p.add_run(h); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= center_from else WD_ALIGN_PARAGRAPH.LEFT

    # data rows
    for ri, row in enumerate(rows):
        is_tot = any('TOTAL' in str(v).upper() or str(v).strip() in ('∑',) for v in row)
        bg = HEX_TOTAL if is_tot else (HEX_EVEN if ri%2==0 else HEX_ODD)
        tr = t.rows[ri+1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg); set_cell_pad(c)
            p = c.paragraphs[0]; p.clear()
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if is_tot: r.bold = True; r.font.color.rgb = NAVY
            else: r.font.color.rgb = DARK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= center_from else WD_ALIGN_PARAGRAPH.LEFT

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows: row.cells[ci].width = Cm(w)

    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Cm(2.8)
    sec.top_margin = sec.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK

# ── PAGE DE COUVERTURE ───────────────────────────────────────────
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(cover)
c = cover.rows[0].cells[0]
set_cell_bg(c, HEX_NAVY)
set_cell_pad(c, top=350, bot=280, left=600, right=600)
c.width = Cm(16)

pt = c.paragraphs[0]
pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = pt.add_run("BIOGUARD ACCESS\n")
r1.bold = True; r1.font.size = Pt(38); r1.font.color.rgb = WHITE; r1.font.name = 'Calibri'
r2 = pt.add_run("Système IoT de Contrôle d'Accès Biométrique")
r2.italic = True; r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xF0); r2.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

pt2 = doc.add_paragraph()
pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(pt2, before=8, after=2)
r3 = pt2.add_run("DOCUMENTATION DU PROJET IoT")
r3.bold = True; r3.font.size = Pt(20); r3.font.color.rgb = NAVY

pt3 = doc.add_paragraph()
pt3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(pt3, before=0, after=30)
r4 = pt3.add_run("Défi Dragons' Den IoT 2026  —  Session accélérée H2026")
r4.font.size = Pt(13); r4.font.color.rgb = BLUE

add_hr(doc)

ppres = doc.add_paragraph()
ppres.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(ppres, before=18, after=4)
rp = ppres.add_run("Présenté par")
rp.bold = True; rp.font.size = Pt(11); rp.font.color.rgb = GRAY

for name in ["Tiavina Misandratra Andrianantenaina", "Noel Laurian Simen"]:
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(pn, before=3, after=3)
    rn = pn.add_run(name)
    rn.bold = True; rn.font.size = Pt(13); rn.font.color.rgb = NAVY

add_hr(doc)

pdate = doc.add_paragraph()
pdate.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(pdate, before=14, after=0)
rd = pdate.add_run("Mars 2026")
rd.font.size = Pt(12); rd.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION 1 — CONTEXTE
# ════════════════════════════════════════════════════════════════════
section_title(doc, "1", "Contexte du projet")

body(doc,
     "Le projet BioGuard Access s'inscrit dans le cadre d'un projet académique réalisé durant la session "
     "accélérée H2026. Il consiste en la conception et le développement d'un prototype IoT de contrôle "
     "d'accès biométrique intelligent, combinant composantes matérielles et logicielles.")
body(doc,
     "Dans un contexte où la sécurité des accès physiques devient de plus en plus importante — notamment "
     "dans les environnements résidentiels et professionnels — ce projet propose une solution moderne basée "
     "sur l'intégration de technologies embarquées et mobiles. Contrairement aux systèmes traditionnels "
     "(cartes, codes PIN), BioGuard Access utilise la biométrie pour garantir une authentification plus "
     "sécurisée et personnalisée.")
body(doc,
     "Le système repose sur une architecture hybride : un Raspberry Pi assure la capture biométrique et "
     "la prise de décision locale ; une application mobile Expo (React Native) permet la supervision et "
     "l'administration ; une communication en temps réel s'effectue via MQTT ; Firebase gère "
     "l'authentification et la synchronisation cloud.")
body(doc,
     "Ce projet met également l'accent sur la collaboration entre les membres de l'équipe, chacun étant "
     "responsable de modules spécifiques : intégration matérielle, backend IoT, expérience utilisateur mobile.")

# ════════════════════════════════════════════════════════════════════
# SECTION 2 — DESCRIPTION
# ════════════════════════════════════════════════════════════════════
section_title(doc, "2", "Description du projet")

body(doc,
     "BioGuard Access est un prototype fonctionnel de contrôle d'accès intelligent basé sur "
     "l'authentification biométrique. Il vise à offrir une solution sécurisée, fiable et connectée pour "
     "la gestion des accès physiques.")
body(doc,
     "Le Raspberry Pi joue un rôle central dans la capture et la validation des données biométriques, "
     "tandis que l'application mobile permet aux administrateurs de superviser les accès, gérer les "
     "utilisateurs autorisés et consulter les informations en temps réel. Le protocole MQTT assure une "
     "communication efficace entre composants ; Firebase gère la sécurité des comptes et la synchronisation.")

body(doc, "Le système permet :", bold=False)
for item in [
    "L'enrôlement des utilisateurs",
    "L'identification biométrique",
    "La gestion des accès",
    "La consultation des historiques",
    "La supervision en temps réel"
]:
    bullet(doc, item)

# ════════════════════════════════════════════════════════════════════
# SECTION 3 — PROBLÈME
# ════════════════════════════════════════════════════════════════════
section_title(doc, "3", "Problème Ciblé")
sub_title(doc, "Quel problème réel résolvez-vous ?")

body(doc,
     "Les petites et moyennes structures — entreprises, laboratoires, écoles — ont besoin d'un contrôle "
     "d'accès plus fiable qu'un simple code PIN ou une carte magnétique. Pourtant, les solutions "
     "biométriques industrielles restent inaccessibles : coûteuses, complexes à déployer et dépendantes "
     "de serveurs centraux.")

body(doc, "Les problèmes concrets sont :", color=NAVY)
for item in [
    "Le partage ou la divulgation d'un code PIN compromet immédiatement la sécurité.",
    "L'absence de traçabilité empêche tout audit en cas d'incident.",
    "Les solutions professionnelles existantes coûtent plusieurs milliers de dollars.",
    "Les systèmes bon marché n'offrent ni supervision distante ni résistance aux pannes réseau."
]:
    bullet(doc, item)

body(doc,
     "VeinGuard résout ce problème avec un système biométrique embarqué, abordable et supervisable depuis "
     "un smartphone — sans serveur central coûteux. La reconnaissance repose sur la paume et la géométrie "
     "des doigts via un algorithme PalmCode déterministe exécuté localement sur Raspberry Pi.")

# ════════════════════════════════════════════════════════════════════
# SECTION 4 — CLIENT CIBLE
# ════════════════════════════════════════════════════════════════════
section_title(doc, "4", "Client Cible")
sub_title(doc, "Qui achèterait notre produit ?")

table(doc,
    headers=["Segment", "Profil", "Besoin principal"],
    rows=[
        ["PME & Bureaux",       "Locaux sensibles : salle serveur, archives, réserve",                        "Traçabilité + supervision mobile + aucun prestataire"],
        ["Secteur Éducatif",    "Cégeps et écoles enseignant l'IoT, Python, la sécurité embarquée",           "Kit pédagogique tout-en-un (Pi, MQTT, biométrie, Firebase)"],
        ["Makerspaces",         "Espaces partagés gérant l'accès à du matériel de valeur",                    "Déploiement simple, historique, app mobile intuitive"],
        ["Projets académiques", "Universités et cégeps cherchant un projet IoT de bout en bout démontrable",  "Architecture claire, biométrie explicable"],
    ],
    col_widths=[3.5, 6.5, 6.0],
    center_from=10,
)

# ════════════════════════════════════════════════════════════════════
# SECTION 5 — COÛT DE PRODUCTION
# ════════════════════════════════════════════════════════════════════
section_title(doc, "5", "Coût de Production Unitaire")

table(doc,
    headers=["Composant", "Prototype", "Série estimée"],
    rows=[
        ["Raspberry Pi 5 (unité centrale + caméra)",        "165,90 $",      " 65,00 $"],
        ["Capteur de lumière + LCD I2C 16×2 + Résistances", "  9,40 $",      "  8,00 $"],
        ["LEDs, résistances, buzzer, câblage, Breadboard",  " 24,70 $",      "  5,00 $"],
        ["Alimentation (bloc 5 V / 5 A)",                   " 16,95 $",      "  8,00 $"],
        ["Boîtier (impression 3D → moule injection)",        " 12,45 $",      " 12,00 $"],
        ["TOTAL COÛT UNITAIRE",                             "229,15 $ CAD",  " 98,00 $ CAD"],
    ],
    col_widths=[9.0, 3.0, 3.0],
    center_from=1,
)

body(doc,
     "Note : le coût prototype reflète les prix au détail unitaire. En production série, les achats en volume "
     "permettent de réduire le coût unitaire de façon significative.",
     italic=True, color=GRAY)

# ════════════════════════════════════════════════════════════════════
# SECTION 6 — PRIX DE VENTE
# ════════════════════════════════════════════════════════════════════
section_title(doc, "6", "Prix de Vente Proposé")

big_val(doc, "Prix de vente", "499,99 $ – 699,99 $ CAD",
        sub_text="Positionnement haut de gamme accessible — sous la barre psychologique des 700 $")

sub_title(doc, "Prix psychologique")
body(doc,
     "Le seuil de 499,99 $ positionne VeinGuard comme un équipement sérieux de sécurité embarquée, "
     "sans atteindre le territoire des solutions industrielles à plusieurs milliers de dollars. "
     "Ce prix est cohérent avec un kit IoT académique premium.")

sub_title(doc, "Positionnement marché")
body(doc,
     "VeinGuard n'est pas un gadget — c'est un système de contrôle d'accès complet et opérationnel. "
     "Il intègre un algorithme biométrique PalmCode propriétaire, une caméra Raspberry Pi, un protocole "
     "MQTT temps réel, une persistance SQLite locale résistante aux pannes réseau, une synchronisation "
     "Firebase et une application mobile iOS & Android.")

sub_title(doc, "Comparatif concurrents")
table(doc,
    headers=["Produit", "Prix", "Biométrie", "App mobile", "Supervision distante"],
    rows=[
        ["VeinGuard",               "499,99 $", "Oui (paume)", "Oui", "Oui"],
        ["HID iCLASS SE",           "~800 $",   "Non",         "Non", "Partielle"],
        ["ZKTeco MB10VL",           "~350 $",   "Veines main", "Non", "Non"],
        ["Raspberry Pi + badge DIY","~60 $",    "Non",         "Non", "Non"],
    ],
    col_widths=[4.5, 2.2, 3.0, 2.5, 3.8],
    center_from=1,
)

# ════════════════════════════════════════════════════════════════════
# SECTION 7 — MARGE DE PROFIT
# ════════════════════════════════════════════════════════════════════
section_title(doc, "7", "Marge de Profit")
sub_title(doc, "Communication MQTT (API principale)")

table(doc,
    headers=["Indicateur", "Phase prototype", "Production série"],
    rows=[
        ["Prix de vente",                   "499,99 $",  "499,99 $"],
        ["Coût de production unitaire",     "229,15 $",  " 98,00 $"],
        ["Profit unitaire",                 "270,84 $",  "401,99 $"],
        ["Marge brute",                     "   54 %",   "    80 %"],
        ["Projection — 500 unités (série)", "—",         "200 995,00 $ CAD"],
    ],
    col_widths=[7.5, 3.5, 5.0],
    center_from=1,
)

# ════════════════════════════════════════════════════════════════════
# SECTION 8 — INVESTISSEMENT
# ════════════════════════════════════════════════════════════════════
section_title(doc, "8", "Investissement Demandé au Dragon")
sub_title(doc, "Combien demandez-vous ? Pour quel pourcentage ?")

big_val(doc, "Investissement demandé", "25 000 $",
        sub_text="pour 15 % des parts  —  Valorisation implicite : 166 666 $ CAD")

sub_title(doc, "À quoi servira l'investissement ?")
table(doc,
    headers=["#", "Poste budgétaire", "Montant", "Résultat visé"],
    rows=[
        ["1", "Achat en volume — Raspberry Pi 5 (contrats distributeurs)", "10 000 $", "< 65 $ / unité"],
        ["2", "Design industriel — boîtier moulé injection + certifications","8 000 $",  "Finition commerciale"],
        ["3", "App VeinGuard — notifications push + mode hors-ligne complet","5 000 $",  "iOS & Android v1.0"],
        ["4", "Marketing & distribution — 1re série (PME + cégeps)",         "2 000 $",  "Lancement 1re série"],
        ["∑", "TOTAL INVESTISSEMENT",                                        "25 000 $", "15 % des parts"],
    ],
    col_widths=[0.8, 8.2, 2.5, 4.5],
    center_from=0,
)

# ════════════════════════════════════════════════════════════════════
# SECTION 9 — CONCLUSION
# ════════════════════════════════════════════════════════════════════
section_title(doc, "9", "Conclusion")

body(doc,
     "Le projet BioGuard Access constitue un prototype complet intégrant : matériel embarqué, "
     "communication IoT, application mobile, stockage cloud et biométrie. Il permet de démontrer "
     "une solution cohérente, fonctionnelle et réaliste de contrôle d'accès intelligent, tout en "
     "répondant aux objectifs pédagogiques du projet académique.")
body(doc,
     "Sa principale force n'est pas d'être un produit industriel fini, mais d'être un système de bout "
     "en bout : compréhensible, démontrable, maintenable et suffisamment réaliste pour une présentation "
     "technique convaincante.")

# ═══ PIED DE PAGE ════════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("BioGuard Access — Documentation confidentielle  |  Tiavina & Noel, Mars 2026  |  Page ")
    fr.font.size = Pt(8); fr.italic = True; fr.font.color.rgb = MID_GRAY
    add_page_number(fp)

# ─── SAUVEGARDE ─────────────────────────────────────────────────
doc.save("DocumentationIOT.docx")
print("✅  DocumentationIOT.docx mis en forme avec succès !")
