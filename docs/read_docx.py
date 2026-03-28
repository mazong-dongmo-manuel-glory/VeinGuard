from docx import Document
import json

doc = Document("DocumentationIOT.docx")
data = []
for para in doc.paragraphs:
    style = para.style.name
    text  = para.text
    if text.strip():
        data.append({"style": style, "text": text})

# tables
for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    data.append({"style": f"TABLE_{ti}", "rows": rows})

with open("docx_content.json", "w", encoding="utf-8") as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"✅ {len(data)} blocs extraits → docx_content.json")
